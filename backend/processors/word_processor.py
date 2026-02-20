from docx import Document
from typing import Any, Dict
from .base_processor import BaseProcessor


class WordProcessor(BaseProcessor):
    """Processor for modern Word documents (.docx)"""

    def extract_text(self, file_path_or_content: Any) -> str:
        """
        Extract text content from .docx file

        Args:
            file_path_or_content: File path or file-like object

        Returns:
            Extracted text as string
        """
        try:
            doc = Document(file_path_or_content)
            text_parts = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))

            return '\n'.join(text_parts).strip()

        except Exception as e:
            raise Exception(f"Failed to extract text from .docx file: {str(e)}")

    def get_file_type(self) -> str:
        """Get the file type this processor handles"""
        return "docx"

    def get_metadata(self, file_path_or_content: Any) -> Dict[str, Any]:
        """
        Extract metadata from .docx file

        Args:
            file_path_or_content: File path or file-like object

        Returns:
            Dictionary containing document metadata
        """
        metadata = {
            "file_type": "docx"
        }

        try:
            doc = Document(file_path_or_content)

            # Count elements
            metadata["num_paragraphs"] = len(doc.paragraphs)
            metadata["num_tables"] = len(doc.tables)

            # Extract core properties
            core_properties = doc.core_properties
            if core_properties.title:
                metadata["title"] = core_properties.title
            if core_properties.author:
                metadata["author"] = core_properties.author
            if core_properties.subject:
                metadata["subject"] = core_properties.subject
            if core_properties.keywords:
                metadata["keywords"] = core_properties.keywords
            if core_properties.created:
                metadata["created"] = core_properties.created.isoformat()
            if core_properties.modified:
                metadata["modified"] = core_properties.modified.isoformat()

            # Count words
            text = self.extract_text(file_path_or_content)
            metadata["word_count"] = len(text.split())
            metadata["char_count"] = len(text)

        except Exception as e:
            metadata["error"] = str(e)

        return metadata
